import streamlit as st
import ffmpeg
import speech_recognition as sr
from googletrans import Translator
from gtts import gTTS
from docx import Document
import yt_dlp
import os
from docx import Document
from docx.shared import Pt

# Function to extract audio from video
def extract_audio(video_path):
    audio_path = "extracted_audio.wav"
    (
        ffmpeg
        .input(video_path)
        .output(audio_path, acodec="pcm_s16le")
        .run(overwrite_output=True)
    )
    return audio_path

# Function to transcribe audio to text
def transcribe_audio(audio_path):
    recognizer = sr.Recognizer()
    with sr.AudioFile(audio_path) as source:
        audio = recognizer.record(source)
    return recognizer.recognize_google(audio)

# Function to translate text
def translate_text(text, target_lang, translator, chunk_size=200):
    sentences = text.split(". ")  # Split into sentences
    translated_sentences = []

    for i in range(0, len(sentences), chunk_size):
        chunk = ". ".join(sentences[i:i+chunk_size])
        translated_chunk = translator.translate(chunk, dest=target_lang).text
        translated_sentences.append(translated_chunk)

    return " ".join(translated_sentences)

# Function to generate a voiceover
def generate_voiceover(translated_text, language):
    tts = gTTS(text=translated_text, lang=language)
    voiceover_path = f"voiceover_{language}.mp3"
    tts.save(voiceover_path)
    return voiceover_path

# Function to create subtitle file (SRT)
def create_srt(translated_text, language, video_duration):
    srt_path = f"subtitles_{language}.srt"
    with open(srt_path, "w", encoding="utf-8") as f:
        start_time = "00:00:00,000"
        end_time = f"00:{int(video_duration // 60):02}:{int(video_duration % 60):02},000"  # Full duration
        f.write(f"1\n{start_time} --> {end_time}\n{translated_text.strip()}\n\n")
    return srt_path

# Function to replace original audio with translated voiceover
def replace_audio_with_translation(video_path, voiceover_path, lang):
    output_path = f"final_video_{lang}.mp4"

    try:
        # Load video without original audio and add new translated voiceover
        video = ffmpeg.input(video_path).video
        voiceover = ffmpeg.input(voiceover_path).audio

        (
            ffmpeg
            .output(video, voiceover, output_path, vcodec="libx264", acodec="aac", strict='experimental')
            .run(overwrite_output=True)
        )

    except Exception as e:
        st.error(f"Error merging video with translated audio: {e}")
        return None

    return output_path

# Function to download a YouTube video
def download_youtube_video(url):
    output_file = "downloaded_video.mp4"

    # Delete existing file if it exists
    if os.path.exists(output_file):
        os.remove(output_file)

    ydl_opts = {
        'format': 'bestvideo+bestaudio/best',
        'outtmpl': output_file,  # Save as this filename
        'merge_output_format': 'mp4',
        'noplaylist': True
    }

    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        ydl.download([url])

    return output_file

# Function to create a translation document
def create_translation_document(original_text, translations, languages):
    doc = Document()
    doc.add_heading('Translations', level=1)

    # Add Original Text
    doc.add_heading('Original Text:', level=2)
    paragraph = doc.add_paragraph(original_text)
    paragraph.runs[0].font.name = "Nirmala UI"  # Use a font that supports Indian languages
    paragraph.runs[0].font.size = Pt(12)

    # Add Translations
    for lang, translation in zip(languages, translations):
        doc.add_heading(f'Translation in {lang}:', level=2)
        para = doc.add_paragraph(translation)
        para.runs[0].font.name = "Nirmala UI"  # Ensuring proper font rendering
        para.runs[0].font.size = Pt(12)

    # Save the document
    doc_path = "translations.docx"
    doc.save(doc_path)
    return doc_path


# Function to get video duration
def get_video_duration(video_path):
    probe = ffmpeg.probe(video_path)
    return float(probe['format']['duration'])  # Duration in seconds

# Main processing function
def main(video_path, target_languages):
    translator = Translator()
    audio_path = extract_audio(video_path)
    transcribed_text = transcribe_audio(audio_path)
    video_duration = get_video_duration(video_path)
    translations = []

    for lang_code in target_languages:
        translated_text = translate_text(transcribed_text, lang_code, translator)
        translations.append(translated_text)

        # Generate translated voiceover
        voiceover_path = generate_voiceover(translated_text, lang_code)

        # Create subtitles file
        subtitle_path = create_srt(translated_text, lang_code, video_duration)

        # Replace video audio with translated voiceover
        translated_video_path = replace_audio_with_translation(video_path, voiceover_path, lang_code)

        if translated_video_path:
            st.success(f"Video translated successfully for language: {lang_code}")

    # Save all translations in a document
    doc_path = create_translation_document(transcribed_text, translations, target_languages)

    return translations, doc_path

# Streamlit UI
st.title("Software for dubbing of videos")

# Upload video OR enter a YouTube link
video_file = st.file_uploader("Upload a video file", type=["mp4", "mov"])
video_url = st.text_input("Or enter a YouTube URL:")

# Select target languages
target_languages = st.multiselect("Select target languages", ["hi", "te", "ta", "kn", "mr", "ur", "ml", "pa", "gu"])

if st.button("Translate & Subtitle"):
    if video_file:
        video_path = "uploaded_video.mp4"
        with open(video_path, "wb") as f:
            f.write(video_file.getbuffer())
    elif video_url:
        st.info("Downloading YouTube video...")
        video_path = download_youtube_video(video_url)
        st.success("YouTube video downloaded!")
    else:
        st.warning("Please upload a video or enter a YouTube URL.")
        st.stop()

    if target_languages:
        translations, doc_path = main(video_path, target_languages)
        st.success("Processing Completed!")

        # Provide download buttons for translated videos
        for lang in target_languages:
            video_file_to_download = f"final_video_{lang}.mp4"
            if os.path.exists(video_file_to_download):
                with open(video_file_to_download, "rb") as video_file:
                    st.download_button(label=f"Download {lang} Video", data=video_file, file_name=f"{lang}_translated_video.mp4")
            else:
                st.warning(f"No video file found for language: {lang}")

        # Provide download button for translations document
        with open(doc_path, "rb") as doc_file:
            st.download_button(label="Download Translation Document", data=doc_file, file_name="translations.docx")

    else:
        st.warning("Please select at least one language.")
